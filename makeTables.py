from docx.api import Document
import mysql.connector


document = Document('table.docx')

table_names = ["Depositor_Table","Item_Table","Customer_Table","Salesman_Table","Sales_Order_Table","Branch_Table","Customer_new_Table","Loan_Table","Borrower_Table","Account_Table","Depositor_new_Table"]
	
for i in range(len(document.tables)):

	table = document.tables[i]
	data = []
	keys = None
	for j, row in enumerate(table.rows):
	    text = (cell.text.encode('utf-8') for cell in row.cells)
	    if j == 0:
		keys = tuple(text)
		continue
	    row_data = tuple(text)
	    data.append(row_data)
	    
	mydb = mysql.connector.connect(
	  host="localhost",
	  user="francesco",
	  passwd="some_pass",
	  database="test"
	)
	cursor = mydb.cursor()
	sql_item ="CREATE TABLE " + table_names[i] +  " ("
	
	for j in range(len(keys)):
		header = keys[j]
		repeat = str(header).strip().replace(" ","").replace("\n","") + " CHAR(20) NOT NULL,"
		if j == (len(keys)-1):
			repeat = str(header) + " CHAR(20) NOT NULL)"
		sql_item += repeat

	#print(sql_item)	
	cursor.execute(sql_item)
	insert_sql_header = "INSERT INTO " + table_names[i] + " ("
	sql_item = ""
	for j in range(len(keys)):
		header = keys[j]
		repeat = str(header).strip().replace(" ","").replace("\n","") + ","
		if j == (len(keys)-1):
			repeat = str(header)
		sql_item += repeat

	sql_header = ") VALUES ("
	sql_repeat = ""

	if len(keys)==1:
		sql_repeat = "%s)"

	else:
		for j in range(len(keys)-1):
			sql_repeat+="%s, "
		 
	 	sql_repeat += "%s)"

	insert_sql = insert_sql_header + sql_item + sql_header + sql_repeat 
	#print insert_sql
	for j in range(len(keys)):
		val = data[j]
		cursor.execute(insert_sql, val)

			
	mydb.commit()

	print "Table " + str(i) + " complete..." 

	



	
 


